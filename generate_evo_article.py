"""
Generate a professional Word document article about EVO Protocol.
Designed for degens — simple language, lots of visuals.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image, ImageDraw, ImageFont
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "EVO_Article.docx")
IMG_DIR = os.path.join(os.path.dirname(__file__), "_article_images")
os.makedirs(IMG_DIR, exist_ok=True)

# ── Colors ──
C_BG_DARK = "#0A0B0F"
C_SURFACE = "#16181F"
C_SURFACE2 = "#1E2129"
C_BORDER = "#2A2D38"
C_TEXT = "#E4E6EB"
C_TEXT_STRONG = "#FFFFFF"
C_MUTED = "#8A8F9A"
C_DIM = "#5A5F6A"
C_ACCENT = "#818CF8"  # indigo
C_ACCENT_SOFT = "#818CF820"
C_POSITIVE = "#22C55E"
C_NEGATIVE = "#EF4444"
C_WARN = "#F59E0B"
C_PURPLE = "#A855F7"

def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, color_hex):
    """Set table cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with dict of sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            tag = f'w:{edge}'
            elem = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{kwargs[edge].get("sz", "4")}" '
                f'w:color="{kwargs[edge].get("color", "000000")}"/>'
            )
            tcPr.append(elem)

def add_shaded_paragraph(doc, text, bg_color, text_color=None, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """Add a paragraph with background shading."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.space_after = Pt(space_after)
    # Use a single-cell table for background shading
    return p

def section_header(doc, text, color=C_ACCENT):
    """Add a section header with colored bar."""
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    p.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
    run.font.name = 'Calibri'
    return p

def sub_header(doc, text, color=C_ACCENT):
    """Add a sub-section header."""
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    run.font.name = 'Calibri'
    return p

def body_text(doc, text, bold=False, size=11, color=C_TEXT, space_after=8):
    """Add body text."""
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    if bold:
        run.bold = True
    return p

def mixed_text(doc, parts, size=11, space_after=8):
    """Add paragraph with mixed formatting. parts = [(text, bold, color_hex), ...]"""
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    for text, bold, color in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))
        run.bold = bold
    return p

def add_callout_box(doc, title, body_lines, border_color=C_ACCENT, bg_color=C_ACCENT_SOFT):
    """Add a shaded callout box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_cell_border(cell, 
        top={"sz": "12", "color": border_color.lstrip('#')},
        bottom={"sz": "12", "color": border_color.lstrip('#')},
        left={"sz": "12", "color": border_color.lstrip('#')},
        right={"sz": "12", "color": border_color.lstrip('#')},
    )
    # Title
    p = cell.paragraphs[0]
    p.space_before = Pt(8)
    p.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*hex_to_rgb(border_color))
    # Body lines
    for line in body_lines:
        p = cell.add_paragraph()
        p.space_after = Pt(4)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT))
    # Padding
    doc.add_paragraph().space_after = Pt(4)

# ── Generate diagram images with PIL ──

def make_nft_vs_evo_diagram():
    """Create a visual comparison: NFT vs EVO"""
    w, h = 800, 400
    img = Image.new('RGB', (w, h), hex_to_rgb(C_BG_DARK))
    draw = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 18)
        font_body = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 13)
        font_small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 11)
        font_bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 14)
    except:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_bold = ImageFont.load_default()
    
    # Left side: NFT
    x = 50
    draw.rounded_rectangle([x, 60, x+320, 360], radius=12, outline=hex_to_rgb(C_BORDER), width=2)
    draw.text((x+20, 80), "Traditional NFT", fill=hex_to_rgb(C_DIM), font=font_bold)
    
    # NFT image placeholder
    draw.rounded_rectangle([x+100, 120, x+220, 220], radius=8, outline=hex_to_rgb(C_BORDER), width=2)
    draw.text((x+130, 160), "IMG", fill=hex_to_rgb(C_DIM), font=font_body)
    
    draw.text((x+20, 240), "image file on IPFS", fill=hex_to_rgb(C_TEXT), font=font_body)
    draw.text((x+20, 262), "price = whatever someone pays", fill=hex_to_rgb(C_MUTED), font=font_small)
    draw.text((x+20, 284), "floor = 0 (can go to nothing)", fill=hex_to_rgb(C_NEGATIVE), font=font_small)
    draw.text((x+20, 306), "static image forever", fill=hex_to_rgb(C_DIM), font=font_small)
    
    # Right side: EVO
    x2 = 430
    draw.rounded_rectangle([x2, 60, x2+320, 360], radius=12, outline=hex_to_rgb(C_ACCENT), width=3)
    draw.text((x2+20, 80), "EVO Asset", fill=hex_to_rgb(C_ACCENT), font=font_bold)
    
    # Lock icon
    cx, cy = x2+160, 170
    draw.rounded_rectangle([cx-30, cy-15, cx+30, cy+25], radius=4, outline=hex_to_rgb(C_ACCENT), width=2)
    draw.arc([cx-18, cy-40, cx+18, cy-10], start=180, end=360, fill=hex_to_rgb(C_ACCENT), width=2)
    draw.text((cx-22, cy-5), "SOL", fill=hex_to_rgb(C_ACCENT), font=font_small)
    draw.text((x2+110, 205), "locked in PDA", fill=hex_to_rgb(C_ACCENT), font=font_small)
    
    draw.text((x2+20, 240), "SOL inside PDA", fill=hex_to_rgb(C_TEXT), font=font_body)
    draw.text((x2+20, 262), "price = market + locked floor", fill=hex_to_rgb(C_MUTED), font=font_small)
    draw.text((x2+20, 284), "floor = locked SOL (never zero)", fill=hex_to_rgb(C_POSITIVE), font=font_small)
    draw.text((x2+20, 306), "evolving visual state", fill=hex_to_rgb(C_DIM), font=font_small)
    
    path = os.path.join(IMG_DIR, "nft_vs_evo.png")
    img.save(path)
    return path

def make_lifecycle_diagram():
    """Create lifecycle flow diagram"""
    w, h = 800, 350
    img = Image.new('RGB', (w, h), hex_to_rgb(C_BG_DARK))
    draw = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 14)
        font_small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 10)
        font_bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 12)
    except:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_bold = ImageFont.load_default()
    
    boxes = [
        ("FORGE", "lock SOL\nmint asset", C_ACCENT, 50),
        ("FEED", "permissionless\nadvance state", C_PURPLE, 220),
        ("EVOLVE", "visual change\non-chain", C_WARN, 390),
        ("TRADE", "list / buy\nmarketplace", C_POSITIVE, 560),
    ]
    
    for i, (label, desc, color, x) in enumerate(boxes):
        # Box
        draw.rounded_rectangle([x, 80, x+130, 180], radius=10, outline=hex_to_rgb(color), width=2)
        draw.text((x+30, 100), label, fill=hex_to_rgb(color), font=font_bold)
        lines = desc.split('\n')
        for j, line in enumerate(lines):
            draw.text((x+10, 125+j*15), line, fill=hex_to_rgb(C_MUTED), font=font_small)
        
        # Arrow to next
        if i < len(boxes) - 1:
            ax = x + 130
            draw.line([ax+5, 130, ax+35, 130], fill=hex_to_rgb(C_BORDER), width=2)
            draw.polygon([(ax+35, 125), (ax+40, 130), (ax+35, 135)], fill=hex_to_rgb(C_BORDER))
    
    # Shatter box (below, pointing from all)
    sx = 300
    draw.rounded_rectangle([sx, 230, sx+200, 310], radius=10, outline=hex_to_rgb(C_NEGATIVE), width=2)
    draw.text((sx+50, 245), "SHATTER", fill=hex_to_rgb(C_NEGATIVE), font=font_bold)
    draw.text((sx+20, 270), "burn asset, reclaim SOL", fill=hex_to_rgb(C_MUTED), font=font_small)
    
    # Arrows from trade and forge to shatter
    draw.line([115, 180, 115, 250], fill=hex_to_rgb(C_DIM), width=1)
    draw.line([115, 250, 300, 250], fill=hex_to_rgb(C_DIM), width=1)
    draw.line([625, 180, 625, 250], fill=hex_to_rgb(C_DIM), width=1)
    draw.line([625, 250, 500, 250], fill=hex_to_rgb(C_DIM), width=1)
    draw.text((350, 200), "exit always available", fill=hex_to_rgb(C_DIM), font=font_small)
    
    path = os.path.join(IMG_DIR, "lifecycle.png")
    img.save(path)
    return path

def make_architecture_diagram():
    """Create architecture stack diagram"""
    w, h = 800, 420
    img = Image.new('RGB', (w, h), hex_to_rgb(C_BG_DARK))
    draw = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 14)
        font_small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 10)
        font_bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 12)
    except:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_bold = ImageFont.load_default()
    
    layers = [
        ("You (Browser + Wallet)", "Phantom / Solflare signs transactions", C_DIM, 30),
        ("Meld Terminal (Frontend)", "Next.js | wallet adapter | RPC calls | visual rendering", C_ACCENT, 110),
        ("EVO Protocol (On-Chain Program)", "Anchor | 17 instructions | 4 account types | Solana mainnet", C_POSITIVE, 190),
        ("Storage Layer", "Arweave/Irys for images | Merkle-verified manifests | Supabase for logos", C_WARN, 270),
        ("Solana Mainnet", "Program ID: Aw4mAC5... | PDAs hold locked SOL | atomic settlement", C_PURPLE, 350),
    ]
    
    for label, desc, color, y in layers:
        draw.rounded_rectangle([100, y, 700, y+60], radius=10, outline=hex_to_rgb(color), width=2)
        draw.text((120, y+10), label, fill=hex_to_rgb(color), font=font_bold)
        draw.text((120, y+30), desc, fill=hex_to_rgb(C_MUTED), font=font_small)
        
        # Arrow between layers
        if y < 350:
            draw.line([400, y+60, 400, y+10+80], fill=hex_to_rgb(C_BORDER), width=2)
            # Bidirectional
            draw.polygon([(395, y+62), (400, y+67), (405, y+62)], fill=hex_to_rgb(C_BORDER))
            draw.polygon([(395, y+78+8), (400, y+73+8), (405, y+78+8)], fill=hex_to_rgb(C_BORDER))
    
    path = os.path.join(IMG_DIR, "architecture.png")
    img.save(path)
    return path

def make_floor_value_diagram():
    """Create a diagram showing floor value concept"""
    w, h = 800, 300
    img = Image.new('RGB', (w, h), hex_to_rgb(C_BG_DARK))
    draw = ImageDraw.Draw(img)
    
    try:
        font_bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 14)
        font_small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 11)
        font_title = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 16)
    except:
        font_bold = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_title = ImageFont.load_default()
    
    draw.text((250, 20), "The Floor That Can't Zero", fill=hex_to_rgb(C_TEXT_STRONG), font=font_title)
    
    # NFT price chart (can go to zero)
    draw.text((60, 60), "NFT Price", fill=hex_to_rgb(C_DIM), font=font_bold)
    draw.line([60, 80, 60, 220], fill=hex_to_rgb(C_BORDER), width=2)  # Y axis
    draw.line([60, 220, 350, 220], fill=hex_to_rgb(C_BORDER), width=2)  # X axis
    
    # Zigzag going down to near zero
    points = [(70, 90), (110, 120), (140, 100), (180, 150), (220, 130), (260, 180), (300, 200), (340, 215)]
    for i in range(len(points)-1):
        draw.line([points[i], points[i+1]], fill=hex_to_rgb(C_NEGATIVE), width=2)
    draw.text((280, 225), "floor = 0", fill=hex_to_rgb(C_NEGATIVE), font=font_small)
    
    # EVO price chart (floor supported)
    draw.text((470, 60), "EVO Price", fill=hex_to_rgb(C_ACCENT), font=font_bold)
    draw.line([470, 80, 470, 220], fill=hex_to_rgb(C_BORDER), width=2)
    draw.line([470, 220, 760, 220], fill=hex_to_rgb(C_BORDER), width=2)
    
    # Floor line
    draw.line([470, 180, 760, 180], fill=hex_to_rgb(C_POSITIVE), width=1)
    draw.text((720, 160), "floor = locked SOL", fill=hex_to_rgb(C_POSITIVE), font=font_small)
    
    # Zigzag bouncing off floor
    points2 = [(480, 90), (520, 120), (560, 100), (600, 175), (640, 130), (680, 178), (720, 110), (750, 175)]
    for i in range(len(points2)-1):
        draw.line([points2[i], points2[i+1]], fill=hex_to_rgb(C_ACCENT), width=2)
    draw.text((680, 225), "floor > 0 always", fill=hex_to_rgb(C_POSITIVE), font=font_small)
    
    path = os.path.join(IMG_DIR, "floor_value.png")
    img.save(path)
    return path

def make_thesis_diagram():
    """Create the main thesis / "why it matters" diagram"""
    w, h = 800, 360
    img = Image.new('RGB', (w, h), hex_to_rgb(C_BG_DARK))
    draw = ImageDraw.Draw(img)
    
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 18)
        font_bold = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 13)
        font_small = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 10)
    except:
        font_title = ImageFont.load_default()
        font_bold = ImageFont.load_default()
        font_small = ImageFont.load_default()
    
    # Three pillars
    pillars = [
        ("OWNERSHIP", "like an NFT", "unique, tradeable,\nyours to hold", C_WARN, 60),
        ("LIQUIDITY", "like a token", "shatter = instant\nexit to SOL", C_POSITIVE, 300),
        ("FLOOR VALUE", "neither has", "SOL locked in PDA\ncan't go to zero", C_ACCENT, 540),
    ]
    
    draw.text((180, 20), "EVO = NFT + Token + Floor", fill=hex_to_rgb(C_TEXT_STRONG), font=font_title)
    
    for label, subtitle, desc, color, x in pillars:
        draw.rounded_rectangle([x, 70, x+220, 180], radius=12, outline=hex_to_rgb(color), width=2)
        draw.text((x+50, 85), label, fill=hex_to_rgb(color), font=font_bold)
        draw.text((x+60, 108), subtitle, fill=hex_to_rgb(C_DIM), font=font_small)
        lines = desc.split('\n')
        for j, line in enumerate(lines):
            draw.text((x+20, 130+j*14), line, fill=hex_to_rgb(C_TEXT), font=font_small)
        
        # Plus sign
        if x < 500:
            draw.text((x+235, 115), "+", fill=hex_to_rgb(C_BORDER), font=font_title)
    
    # Bottom result
    draw.rounded_rectangle([150, 210, 650, 310], radius=12, outline=hex_to_rgb(C_ACCENT), width=2)
    draw.text((200, 225), "= Programmable Asset with a Floor", fill=hex_to_rgb(C_ACCENT), font=font_bold)
    draw.text((200, 255), "trade it, evolve it, shatter it - your SOL is always there", fill=hex_to_rgb(C_MUTED), font=font_small)
    draw.text((200, 275), "not a JPEG, not a memecoin, not a bet - a new primitive", fill=hex_to_rgb(C_MUTED), font=font_small)
    
    path = os.path.join(IMG_DIR, "thesis.png")
    img.save(path)
    return path

# ── Build the document ──

def build_document():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT))
    
    # Set document background (dark theme via section properties)
    # Note: Word doesn't easily support full dark background, so we use colored shading on elements
    
    # ── TITLE PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EVO PROTOCOL")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_ACCENT))
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Assets That Don't Stay the Same")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("How programmable assets with a guaranteed floor value")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_MUTED))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("change everything we know about digital ownership")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_MUTED))
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Tech specs box on title page
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    set_cell_bg(cell, C_SURFACE.lstrip('#'))
    set_cell_border(cell,
        top={"sz": "8", "color": C_ACCENT.lstrip('#')},
        bottom={"sz": "8", "color": C_ACCENT.lstrip('#')},
        left={"sz": "8", "color": C_ACCENT.lstrip('#')},
        right={"sz": "8", "color": C_ACCENT.lstrip('#')},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Solana Mainnet  |  Anchor 0.31  |  17 Instructions  |  5 Lifecycle Types")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*hex_to_rgb(C_DIM))
    
    doc.add_page_break()
    
    # ── SECTION 1: THE PROBLEM ──
    section_header(doc, "The Problem: Everything Can Go to Zero")
    
    body_text(doc, "Every digital asset you've ever held has the same fatal flaw. The floor is zero.", size=12)
    
    mixed_text(doc, [
        ("NFTs? ", True, C_WARN),
        ("Beautiful art, strong communities, deep marketplaces. But the floor price? Whatever someone is willing to pay. When hype dies, the floor goes to zero. You're left holding a JPEG worth nothing.", False, C_TEXT),
    ])
    
    mixed_text(doc, [
        ("Meme tokens? ", True, C_NEGATIVE),
        ("Instant liquidity, viral reach, fun as hell. But the LP can be pulled. The devs can rug. The chart can bleed to zero while you sleep. Your tokens are only worth what the pool says they are.", False, C_TEXT),
    ])
    
    mixed_text(doc, [
        ("Gambling? ", True, C_NEGATIVE),
        ("You bet, you watch, you lose. The house always wins. When you lose, you walk away with nothing. There's no asset, no ownership, no exit except winning.", False, C_TEXT),
    ])
    
    body_text(doc, "")
    mixed_text(doc, [
        ("Every single one of these has the same problem: ", False, C_TEXT),
        ("your money disappears and you have nothing to show for it.", True, C_TEXT_STRONG),
    ])
    
    # ── SECTION 2: WHAT IS EVO ──
    section_header(doc, "What is EVO?")
    
    body_text(doc, "EVO is a new on-chain asset primitive on Solana. It's not quite an NFT. It's not quite a token. It's something new.", size=12)
    
    mixed_text(doc, [
        ("Every EVO locks real SOL inside a program-derived address (PDA). ", True, C_ACCENT),
        ("That locked SOL gives each asset a provable floor value that cannot go to zero. The SOL belongs to whoever holds the EVO. Not the creator. Not the protocol. The holder.", False, C_TEXT),
    ])
    
    mixed_text(doc, [
        ("Unlike a static NFT — a frozen image with a market price — an EVO is ", False, C_TEXT),
        ("alive", True, C_ACCENT),
        (". It has lifecycle states. It can reveal, evolve, and change over time based on on-chain triggers. The protocol itself is the source of truth for what the asset looks like at any given moment.", False, C_TEXT),
    ])
    
    mixed_text(doc, [
        ("You can trade EVOs on the built-in marketplace. ", False, C_TEXT),
        ("You can shatter them to reclaim the locked SOL.", True, C_POSITIVE),
        (" You can feed them to advance their state. The exit is always built in.", False, C_TEXT),
    ])
    
    body_text(doc, "")
    
    # NFT vs EVO comparison image
    img1 = make_nft_vs_evo_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img1, width=Inches(6.5))
    
    # ── SECTION 3: THE FLOOR ──
    section_header(doc, "The Floor That Can't Zero")
    
    body_text(doc, "This is the core insight. Everything else flows from this.", size=12)
    
    mixed_text(doc, [
        ("When you forge an EVO, SOL gets locked inside a PDA on Solana. That PDA is controlled by the EVO program, not by any human. The program's rules say: ", False, C_TEXT),
        ("whoever holds the EVO can shatter it and reclaim the SOL.", True, C_ACCENT),
    ])
    
    body_text(doc, "")
    body_text(doc, "This means:")
    
    bullets = [
        ("The floor is real SOL, not a promise.", C_POSITIVE),
        ("The floor can't be rugged — the SOL is in the PDA, not the creator's wallet.", C_POSITIVE),
        ("The exit is always one transaction away — shatter and you get your SOL back.", C_POSITIVE),
        ("The market price can go ABOVE the floor (hype, rarity, evolution) but never BELOW it.", C_ACCENT),
        ("You can lose the premium. You can never lose the floor.", C_TEXT_STRONG),
    ]
    for text, color in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))
        run.font.size = Pt(11)
    
    body_text(doc, "")
    img_floor = make_floor_value_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_floor, width=Inches(6.5))
    
    # ── SECTION 4: HOW IT WORKS ──
    section_header(doc, "How It Works: From Zero to Living Asset")
    
    steps = [
        ("1. Create a Collection", "A creator calls create_collection, paying a small protocol fee (0.0459 SOL). This defines the lifecycle type, randomness policy, metadata, supply ceiling, and fee structure. One-time setup.", C_ACCENT),
        ("2. Forge an EVO", "Users forge (mint) EVOs from the collection. Each forge locks SOL inside a PDA. That SOL is the floor value — it belongs to whoever holds the EVO. The creator sets the lock amount per collection.", C_WARN),
        ("3. It Evolves", "Depending on the lifecycle type, the EVO can reveal, evolve through stages, or change visually. Triggers are on-chain and permissionless where possible. The asset doesn't stay the same — it has a life.", C_PURPLE),
        ("4. Trade It", "List on the built-in marketplace. Buyers pay the listing price. The protocol handles escrow, fees, and ownership transfer atomically. No middleman, no escrow service, no trust needed.", C_POSITIVE),
        ("5. Shatter for Value", "Don't want to sell? Shatter the EVO. It gets burned and the locked SOL is returned to you minus a max 20% fee. The exit is always one transaction away. You're never trapped.", C_NEGATIVE),
        ("6. Verify Everything", "Manifests are Merkle-verified on-chain. Metadata and images are stored on Arweave/Irys. The protocol is the source of truth — not a centralized server, not a dev's database.", C_ACCENT),
    ]
    
    for title, desc, color in steps:
        add_callout_box(doc, title, [desc], border_color=color, bg_color=C_SURFACE)
    
    body_text(doc, "")
    
    # Lifecycle diagram
    img_life = make_lifecycle_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_life, width=Inches(6.5))
    
    # ── SECTION 5: LIFECYCLES ──
    doc.add_page_break()
    section_header(doc, "Lifecycles: Not All EVOs Are the Same")
    
    body_text(doc, "EVO supports five lifecycle types. Creators pick one when deploying a collection. This determines how the asset changes over time.", size=12)
    
    lifecycles = [
        ("Static", "Asset stays the same forever. Classic NFT behavior. Boring but safe.", C_DIM),
        ("Reveal", "Hidden at mint, revealed later by creator or community. Build hype, then drop the reveal.", C_ACCENT),
        ("Commit-Reveal", "Creator commits a hash before mint, reveals after. Proves fairness — no last-second manipulation. Trustless reveals.", C_POSITIVE),
        ("Reveal & Evolve", "Reveals on trigger, then continues evolving through stages. The asset keeps changing after reveal. Living art.", C_WARN),
        ("Custom", "Authority-driven visual stages. Full manual control. The creator decides when the asset changes and what it changes to.", C_NEGATIVE),
    ]
    
    for name, desc, color in lifecycles:
        mixed_text(doc, [
            (f"  {name}:  ", True, color),
            (desc, False, C_TEXT),
        ], space_after=6)
    
    body_text(doc, "")
    mixed_text(doc, [
        ("Three randomness policies ", True, C_ACCENT),
        ("control how reveals work: None (deterministic), Predetermined (creator sets the order), and BatchReveal (on-chain batch randomness for fair distribution).", False, C_TEXT),
    ])
    
    # ── SECTION 6: COMPARISON TABLES ──
    section_header(doc, "The Comparison: NFT vs Meme Token vs EVO")
    
    # Table 1: NFT vs Meme vs EVO
    table = doc.add_table(rows=10, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Property", "NFT", "Meme Token", "EVO"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, C_SURFACE2.lstrip('#'))
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        if i == 0:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
        elif i == 1:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_WARN))
        elif i == 2:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_NEGATIVE))
        else:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_ACCENT))
    
    rows = [
        ("Backed by", "JPEG on IPFS (no SOL)", "Liquidity pool (can drain)", "SOL locked in PDA"),
        ("Floor price", "Can go to zero", "Can go to zero", "Locked SOL (never zero)"),
        ("Fungibility", "Non-fungible (1 of 1)", "Fully fungible", "Non-fungible, shatters to SOL"),
        ("State", "Static (never changes)", "Static (no visual)", "Evolving (on-chain state)"),
        ("Exit", "Sell on marketplace", "Sell into LP", "Shatter (always works)"),
        ("Rug risk", "Creator can abandon", "Devs can pull LP", "No rug — SOL in PDA"),
        ("Speculation", "Hype + scarcity", "Viral + LP depth", "Premium + locked floor"),
        ("Strength", "Deep marketplaces", "Instant liquidity", "Guaranteed floor"),
        ("Weakness", "Floor = 0", "LP can be pulled", "Max 20% shatter fee, new"),
    ]
    
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.cell(ri+1, ci)
            set_cell_bg(cell, C_SURFACE.lstrip('#'))
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
                run.bold = True
            elif ci == 3:
                if "zero" in val.lower() or "always" in val.lower() or "no rug" in val.lower() or "locked" in val.lower() or "guaranteed" in val.lower():
                    run.font.color.rgb = RGBColor(*hex_to_rgb(C_POSITIVE))
                else:
                    run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT))
            elif ci == 1 and ("zero" in val.lower() or "abandon" in val.lower()):
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_NEGATIVE))
            elif ci == 2 and ("zero" in val.lower() or "drain" in val.lower() or "pulled" in val.lower()):
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_NEGATIVE))
            else:
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT))
    
    body_text(doc, "")
    mixed_text(doc, [
        ("NFTs ", True, C_WARN),
        ("built culture and community on art and scarcity. ", False, C_TEXT),
        ("Meme tokens ", True, C_NEGATIVE),
        ("proved liquidity and virality can move real money fast. Both have real value and real risks. ", False, C_TEXT),
        ("EVO ", True, C_ACCENT),
        ("takes the non-fungible ownership of NFTs, the liquidity option of tokens, and adds a guaranteed floor backed by locked SOL.", False, C_TEXT),
    ])
    
    # ── SECTION 7: GAMBLING vs EVO ──
    section_header(doc, "Gambling vs EVO: Same Risk, Different Exit")
    
    body_text(doc, "Both involve risk and SOL. The difference is what you walk away with.", size=12)
    
    table2 = doc.add_table(rows=7, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ["Aspect", "Gambling", "EVO"]
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        set_cell_bg(cell, C_SURFACE2.lstrip('#'))
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        if i == 0:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
        elif i == 1:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_NEGATIVE))
        else:
            run.font.color.rgb = RGBColor(*hex_to_rgb(C_ACCENT))
    
    gambling_rows = [
        ("Your SOL", "Gone the moment you bet", "Locked in PDA you control"),
        ("Floor value", "Zero — lose everything", "Locked SOL minus max 20%"),
        ("Exit", "Only if you win", "Always available via shatter"),
        ("House edge", "Built into the odds", "No house — fee is protocol surcharge"),
        ("Ownership", "Nothing — bet and leave", "Tradeable evolving asset"),
        ("Upside", "Fixed payout if you win", "Market premium + locked floor"),
    ]
    
    for ri, row_data in enumerate(gambling_rows):
        for ci, val in enumerate(row_data):
            cell = table2.cell(ri+1, ci)
            set_cell_bg(cell, C_SURFACE.lstrip('#'))
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT_STRONG))
            elif ci == 1:
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_NEGATIVE))
            else:
                run.font.color.rgb = RGBColor(*hex_to_rgb(C_POSITIVE))
    
    body_text(doc, "")
    mixed_text(doc, [
        ("The key insight: ", True, C_TEXT_STRONG),
        ("gambling is a binary bet with a house that always wins. EVO is an asset you own with a guaranteed exit. The locked SOL is your money, not a wager. The market premium is upside, not a jackpot. You can lose the premium, but you never lose the floor.", False, C_TEXT),
    ])
    
    # ── SECTION 8: THE THESIS ──
    doc.add_page_break()
    section_header(doc, "The Thesis: Why This Changes Everything")
    
    body_text(doc, "EVO sits at the intersection of three things that have never been combined before:", size=12)
    
    img_thesis = make_thesis_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_thesis, width=Inches(6.5))
    
    body_text(doc, "")
    mixed_text(doc, [
        ("NFTs gave us digital ownership. ", True, C_WARN),
        ("Meme tokens gave us instant liquidity. ", True, C_NEGATIVE),
        ("Neither gave us a floor. ", True, C_TEXT_STRONG),
        ("EVO gives us all three.", True, C_ACCENT),
    ])
    
    body_text(doc, "")
    sub_header(doc, "Why this matters for degens:")
    
    degen_points = [
        ("You're never trapped. ", "Hold for the premium, shatter for the floor. The choice is yours, on-chain, permissionless, 24/7."),
        ("No rug possible. ", "The SOL is in a PDA controlled by the program. The creator can't touch it. The protocol enforces the rules. Code is law."),
        ("Evolution = engagement. ", "Assets that change keep people watching. Reveal hype, evolution arcs, stage unlocks — it's not a static JPEG that gets boring after a week."),
        ("Built-in marketplace. ", "No need for Magic Eden integration or third-party escrow. List, buy, settle — all on-chain, all atomic."),
        ("Fair reveals. ", "Commit-reveal with on-chain hashing means creators can't peek and front-run. BatchReveal uses on-chain randomness for trustless distribution."),
        ("Merkle-verified supply. ", "The entire collection manifest is committed as a Merkle root on-chain. Anyone can verify that what you see is what was deployed. No hidden mints, no supply manipulation."),
    ]
    
    for title, desc in degen_points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*hex_to_rgb(C_ACCENT))
        run2 = p.add_run(desc)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(*hex_to_rgb(C_TEXT))
    
    # ── SECTION 9: ARCHITECTURE ──
    section_header(doc, "Under the Hood: Architecture")
    
    body_text(doc, "How the pieces fit together — from your browser to Solana.", size=12)
    
    img_arch = make_architecture_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_arch, width=Inches(6.5))
    
    body_text(doc, "")
    sub_header(doc, "Tech Specs")
    
    specs = [
        ("Network", "Solana mainnet"),
        ("Program ID", "Aw4mAC5oUfQCP65a8a6mTwkrL2CoUMsBa45KvWPY3CN2"),
        ("Framework", "Anchor 0.31.0"),
        ("Instructions", "17"),
        ("Account types", "4 (Protocol, Collection, EVO, Listing)"),
        ("Lifecycle types", "5 (Static, Reveal, CommitReveal, RevealAndEvolve, Custom)"),
        ("Randomness policies", "3 (None, Predetermined, BatchReveal)"),
        ("Max supply per collection", "20,000"),
        ("Collection creation fee", "0.0459 SOL"),
        ("Transfer fee", "0.009 SOL (flat, to treasury)"),
        ("Max shatter fee", "20% (2000 bps)"),
        ("Max royalty", "25% (2500 bps)"),
        ("Storage", "Arweave/Irys for images, Supabase for logos"),
        ("Frontend", "Next.js 16, deployed on Vercel"),
        ("Source code", "github.com/stephenclawdbot-png/EVO"),
    ]
    
    for label, val in specs:
        mixed_text(doc, [
            (f"  {label}:  ", True, C_ACCENT),
            (val, False, C_TEXT),
        ], space_after=4)
    
    # ── SECTION 10: THE 17 INSTRUCTIONS ──
    section_header(doc, "The 17 Instructions")
    
    body_text(doc, "Everything EVO does is one of these 17 on-chain instructions:", size=11)
    
    instructions = [
        ("initialize_protocol", "One-time setup. Sets treasury, authority, and creation fee."),
        ("create_collection", "Deploys a new EVO collection with config + metadata."),
        ("forge", "Mints a new EVO inside a collection. Locks SOL."),
        ("feed", "Permissionless trigger that advances an EVO's lifecycle state."),
        ("evolve", "Advances visual state when lifecycle conditions are met."),
        ("reveal_collection", "Reveals hidden assets using committed secret."),
        ("commit_reveal", "Creator commits hash before mint to prove fairness."),
        ("list", "Lists an EVO on the built-in marketplace."),
        ("delist", "Removes a listing."),
        ("buy", "Purchases a listed EVO. SOL flows to seller, fee to treasury."),
        ("shatter", "Burns the EVO and reclaims locked SOL. Exit is always available."),
        ("transfer", "Sends an EVO to another wallet. Flat fee to treasury."),
        ("set_visual_stage", "Authority-only override for Custom lifecycle."),
        ("update_metadata", "Updates off-chain metadata URI."),
        ("update_treasury", "Rotates treasury address or authority."),
        ("close_collection", "Permanently closes a collection."),
        ("verify_merkle_proof", "Permissionless on-chain manifest verification."),
    ]
    
    for name, desc in instructions:
        mixed_text(doc, [
            (f"  {name}", True, C_ACCENT),
            (f"  —  {desc}", False, C_MUTED),
        ], space_after=3)
    
    # ── CLOSING ──
    doc.add_page_break()
    section_header(doc, "The Bottom Line")
    
    body_text(doc, "EVO is a new primitive. Not a better NFT. Not a better token. A new category.", size=13)
    
    body_text(doc, "")
    mixed_text(doc, [
        ("It takes the best parts of what already works — ", False, C_TEXT),
        ("ownership from NFTs, liquidity from tokens", True, C_WARN),
        (" — and fixes the fatal flaw both share: ", False, C_TEXT),
        ("the floor can go to zero.", True, C_NEGATIVE),
    ])
    
    body_text(doc, "")
    mixed_text(doc, [
        ("By locking real SOL inside each asset, EVO guarantees that every holder has an exit. You're not betting. You're not hoping someone buys your JPEG. You're holding an asset with a provable floor, a market premium on top, and a one-click exit back to SOL.", False, C_TEXT),
    ])
    
    body_text(doc, "")
    mixed_text(doc, [
        ("Assets that don't stay the same. ", True, C_ACCENT),
        ("That's not a tagline. That's the mechanism.", False, C_TEXT_STRONG),
    ])
    
    body_text(doc, "")
    
    # Final callout
    add_callout_box(doc, "EVO Protocol + Meld Terminal", [
        "EVO is the protocol — the on-chain rules that define what an asset is.",
        "Meld is the terminal — where you create, forge, trade, and shatter.",
        "",
        "Live on Solana mainnet.",
        "meldterminal.io",
    ], border_color=C_ACCENT, bg_color=C_SURFACE)
    
    # Save
    doc.save(OUTPUT)
    print(f"Document saved to: {OUTPUT}")

if __name__ == "__main__":
    build_document()